"""
Document writer — creates and combines files in multiple output formats.

create_document(content, format, output_path)    -> Path
combine_documents(paths, format, output_path)    -> Path
convert_document(src_path, target_format)        -> Path

Supported output formats: docx  pdf  html  txt  md  xlsx  csv
"""
from __future__ import annotations

import logging
import re
import textwrap
from pathlib import Path
from typing import Any

from src.automation.document_reader import read_file, read_multiple_files

logger = logging.getLogger(__name__)

_OUT_DIR_NAME = "AutoMotoAI_Documents"


def _output_dir() -> Path:
    d = Path.home() / "Documents" / _OUT_DIR_NAME
    d.mkdir(parents=True, exist_ok=True)
    return d


def _safe_name(name: str) -> str:
    return re.sub(r'[^\w\-. ]', '_', name).strip()


# ──────────────────────────────── writers ─────────────────────────────────────

def _write_txt(content: str, out: Path) -> Path:
    out.write_text(content, encoding="utf-8")
    return out


def _write_html(content: str, out: Path, title: str = "AutoMoto AI Document") -> Path:
    # Convert simple markdown-ish text to HTML
    lines = content.split("\n")
    body_parts: list[str] = []
    in_pre = False
    for line in lines:
        stripped = line.rstrip()
        if stripped.startswith("```") or stripped.startswith("~~~"):
            if in_pre:
                body_parts.append("</pre>")
                in_pre = False
            else:
                body_parts.append('<pre style="background:#f4f4f4;padding:12px;border-radius:6px;overflow-x:auto;">')
                in_pre = True
            continue
        if in_pre:
            body_parts.append(_esc(line))
            continue
        if stripped.startswith("### "):
            body_parts.append(f"<h3>{_esc(stripped[4:])}</h3>")
        elif stripped.startswith("## "):
            body_parts.append(f"<h2>{_esc(stripped[3:])}</h2>")
        elif stripped.startswith("# "):
            body_parts.append(f"<h1>{_esc(stripped[2:])}</h1>")
        elif stripped.startswith("---") or stripped.startswith("==="):
            body_parts.append("<hr>")
        elif not stripped:
            body_parts.append("<br>")
        else:
            # inline bold/code
            line_html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', _esc(stripped))
            line_html = re.sub(r'`(.+?)`', r'<code>\1</code>', line_html)
            body_parts.append(f"<p>{line_html}</p>")
    if in_pre:
        body_parts.append("</pre>")
    body = "\n".join(body_parts)
    html = textwrap.dedent(f"""\
        <!DOCTYPE html>
        <html lang="en">
        <head>
          <meta charset="UTF-8">
          <title>{_esc(title)}</title>
          <style>
            body {{ font-family: "Segoe UI", Arial, sans-serif; max-width: 860px;
                    margin: 40px auto; padding: 0 24px; line-height: 1.6;
                    color: #222; background: #fff; }}
            h1 {{ color: #1a73e8; }} h2 {{ color: #333; }} h3 {{ color: #555; }}
            code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; }}
            hr {{ border: none; border-top: 1px solid #ddd; margin: 20px 0; }}
          </style>
        </head>
        <body>
        {body}
        </body>
        </html>
    """)
    out.write_text(html, encoding="utf-8")
    return out


def _write_docx(content: str, out: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.core_properties.author = "AutoMoto AI"
        for line in content.split("\n"):
            stripped = line.rstrip()
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("---") or stripped.startswith("==="):
                doc.add_paragraph("─" * 60)
            else:
                p = doc.add_paragraph()
                _add_inline_docx(p, stripped)
        doc.save(str(out))
        return out
    except ImportError as exc:
        raise RuntimeError("python-docx not installed: pip install python-docx") from exc


def _add_inline_docx(para, text: str) -> None:
    """Add bold/regular runs to a paragraph based on **bold** markers."""
    from docx.shared import Pt
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def _write_pdf(content: str, out: Path) -> Path:
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Helvetica", size=11)
        for line in content.split("\n"):
            stripped = line.rstrip()
            if stripped.startswith("# "):
                pdf.set_font("Helvetica", "B", 16)
                pdf.cell(0, 10, stripped[2:], ln=True)
                pdf.set_font("Helvetica", size=11)
            elif stripped.startswith("## "):
                pdf.set_font("Helvetica", "B", 13)
                pdf.cell(0, 8, stripped[3:], ln=True)
                pdf.set_font("Helvetica", size=11)
            elif stripped.startswith("### "):
                pdf.set_font("Helvetica", "B", 11)
                pdf.cell(0, 7, stripped[4:], ln=True)
                pdf.set_font("Helvetica", size=11)
            elif stripped.startswith("---") or stripped.startswith("==="):
                pdf.ln(2)
                pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 180, pdf.get_y())
                pdf.ln(2)
            elif not stripped:
                pdf.ln(4)
            else:
                # Remove markdown bold markers for PDF plain text
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
                clean = re.sub(r'`(.+?)`', r'\1', clean)
                # fpdf2 multi_cell handles wrapping
                try:
                    pdf.multi_cell(0, 6, clean)
                except Exception:
                    pdf.multi_cell(0, 6, clean.encode("latin-1", "replace").decode("latin-1"))
        pdf.output(str(out))
        return out
    except ImportError as exc:
        raise RuntimeError("fpdf2 not installed: pip install fpdf2") from exc


def _write_xlsx(content: str, out: Path) -> Path:
    """Write CSV-like tab/comma content into an Excel file."""
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        for row_txt in content.split("\n"):
            if not row_txt.strip():
                continue
            sep = "\t" if "\t" in row_txt else ","
            ws.append(row_txt.split(sep))
        wb.save(str(out))
        return out
    except ImportError as exc:
        raise RuntimeError("openpyxl not installed: pip install openpyxl") from exc


_WRITERS = {
    "txt":  lambda c, o: _write_txt(c, o),
    "md":   lambda c, o: _write_txt(c, o),
    "html": lambda c, o: _write_html(c, o, o.stem),
    "htm":  lambda c, o: _write_html(c, o, o.stem),
    "docx": lambda c, o: _write_docx(c, o),
    "pdf":  lambda c, o: _write_pdf(c, o),
    "xlsx": lambda c, o: _write_xlsx(c, o),
    "csv":  lambda c, o: _write_txt(c, o),
}


# ──────────────────────────────── public API ──────────────────────────────────

def create_document(
    content: str,
    output_format: str,
    output_path: str | Path | None = None,
    filename: str | None = None,
) -> Path:
    """
    Create a new document from *content* in *output_format*.
    Returns the path of the created file.
    """
    fmt = output_format.lstrip(".").lower()
    if fmt not in _WRITERS:
        raise ValueError(f"Unsupported output format '{fmt}'. Supported: {list(_WRITERS)}")

    if output_path:
        out = Path(output_path).expanduser().resolve()
        out.parent.mkdir(parents=True, exist_ok=True)
    else:
        stem = _safe_name(filename or "document")
        out = _output_dir() / f"{stem}.{fmt}"

    result = _WRITERS[fmt](content, out)
    logger.info("Created document: %s", result)
    return result


def combine_documents(
    paths: list[str | Path],
    output_format: str,
    output_path: str | Path | None = None,
    separator: str = "\n\n---\n\n",
) -> Path:
    """
    Read multiple files, merge their text content, and write to output_format.
    Returns the path of the combined file.
    """
    results = read_multiple_files(paths)
    parts: list[str] = []
    for r in results:
        if r.ok:
            parts.append(f"# {r.name}\n\n{r.content}")
        else:
            parts.append(f"# {r.name}\n\n[Could not read: {r.error}]")

    combined = separator.join(parts)

    stems = "_".join(_safe_name(Path(p).stem)[:12] for p in paths[:3])
    if len(paths) > 3:
        stems += f"_and_{len(paths)-3}_more"
    filename = f"combined_{stems}"

    return create_document(combined, output_format, output_path, filename=filename)


def convert_document(
    src_path: str | Path,
    target_format: str,
    output_path: str | Path | None = None,
) -> Path:
    """
    Convert *src_path* from its current format to *target_format*.
    Uses the document reader to extract content, then writes it in the new format.
    """
    result = read_file(src_path)
    if not result.ok:
        raise ValueError(f"Cannot read source file: {result.error}")

    src = Path(src_path)
    filename = src.stem
    return create_document(result.content, target_format, output_path, filename=filename)


def supported_output_formats() -> list[str]:
    return sorted(_WRITERS.keys())


def _esc(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
